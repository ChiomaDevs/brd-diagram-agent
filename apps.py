import streamlit as st
import os
import sys
from dotenv import load_dotenv

# NEW: Import your agent code (adjust path if needed)
import sys
sys.path.insert(0, '.')  # Ensures same-folder imports work
try:
    import enhanced_agent_brd as ag  # Alias to avoid conflicts
except ImportError as e:
    st.error(f"❌ Import failed: {e}. Check enhanced_agent_brd.py exists.")
    st.stop()

# Load env (your API key)
api_key = st.secrets.get("OPENAI_API_KEY")  # Pull from Cloud secrets
if not api_key:
    st.error("❌ OpenAI API key missing in app secrets!")
    st.stop()

# Streamlit Page Config (Title, layout)
st.set_page_config(page_title="BRD Diagram Agent", page_icon="📊", layout="wide")

# Sidebar: Instructions
st.sidebar.title("🤖 BRD Diagram Agent")
st.sidebar.markdown("""
1. Drag-drop a BRD file (PDF, TXT, MD, DOCX) or paste text.
2. Click **Generate Diagrams**.
3. View Mermaid codes + rendered SVGs below.
4. Download images/codes as needed.

**Powered by:** LangChain + OpenAI + Mermaid
""")

# Main Title
st.title("📈 Generate DFD, Logic, and DB Diagrams from Your BRD")
st.markdown("*Upload or paste your Business Requirements Document below.*")

# Input Section
col1, col2 = st.columns([3, 1])  # Split layout: Wide for text/file, narrow for button

with col1:
    # File Uploader (Drag-Drop!)
    uploaded_file = st.file_uploader(
        "Choose a BRD file", 
        type=['pdf', 'txt', 'md', 'docx'],  # UPDATED: Added .docx
        help="Supports PDF, TXT, MD, DOCX. Drag & drop here!"
    )
    
    # Or Paste Text
    brd_text = st.text_area(
        "Or paste BRD text directly:", 
        height=200, 
        placeholder="E.g., Users submit issues via portal. If external, admin allocates..."
    )

with col2:
    # Generate Button
    generate_btn = st.button("🚀 Generate Diagrams", type="primary", use_container_width=True)

# Process Input
if generate_btn or uploaded_file or brd_text.strip():
    if uploaded_file:
        # UPDATED: Enhanced file reading (PDF/TXT/MD/DOCX)
        if uploaded_file.type == "application/pdf":
            try:
                import pypdf2
                reader = pypdf2.PdfReader(uploaded_file)
                brd_text = ""
                for page in reader.pages:
                    brd_text += page.extract_text() + "\n"
                if not brd_text.strip():
                    st.warning("PDF read empty—might be scanned image. Try OCR tools or paste text.")
                else:
                    st.success(f"📄 Loaded PDF: {len(brd_text)} chars")
            except Exception as e:
                st.error(f"PDF read failed: {e}. Paste text instead?")
                brd_text = ""
        elif uploaded_file.name.endswith('.docx'):
            try:
                from docx import Document
                doc = Document(uploaded_file)
                brd_text = "\n".join([para.text for para in doc.paragraphs])
                st.success(f"📄 Loaded DOCX: {len(doc.paragraphs)} paragraphs")
            except Exception as e:
                st.error(f"DOCX read failed: {e}. Install python-docx?")
                brd_text = ""
        else:  # TXT/MD
            brd_text = uploaded_file.read().decode("utf-8")
            st.success(f"📄 Loaded {uploaded_file.name}")
    elif not brd_text.strip():
        st.warning("Add text or upload a file!")
        st.stop()

    # Show Input Preview
    with st.expander("👁️ Preview BRD Text"):
        st.text_area("Preview:", brd_text, height=150, disabled=True)

    # Progress Bar (Fun!)
    progress_bar = st.progress(0)
    status = st.empty()

    # Run Your Agent/Fallback
    status.text("🤖 Parsing BRD...")
    progress_bar.progress(0.3)
    
    if ag.agent is not None:
        # Agent Mode
        result = ag.agent.invoke({"input": brd_text})
        output = result if isinstance(result, str) else result.get("output", str(result))
        status.text("✅ Agent complete!")
    else:
        # Fallback Pipeline
        status.text("🔄 Running pipeline...")
        progress_bar.progress(0.6)
        extracted = ag.parse_brd_tool(brd_text)
        dfd = ag.generate_dfd_tool(extracted)
        logic = ag.generate_logic_tool(extracted)
        db = ag.generate_db_tool(extracted)
        progress_bar.progress(1.0)
        
        # Render & Save (Reuse your function)
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)
        
        dfd_rendered = ag.render_and_save_mermaid(dfd, "DFD", output_dir)
        logic_rendered = ag.render_and_save_mermaid(logic, "Logic", output_dir)
        db_rendered = ag.render_and_save_mermaid(db, "ERD", output_dir)
        
        # NEW: Auto-save .mmd codes to output/
        with open(os.path.join(output_dir, "dfd.mmd"), "w", encoding="utf-8") as f:
            f.write(dfd)
        with open(os.path.join(output_dir, "logic.mmd"), "w", encoding="utf-8") as f:
            f.write(logic)
        with open(os.path.join(output_dir, "erd.mmd"), "w", encoding="utf-8") as f:
            f.write(db)
        
        # NEW: Convert SVG to PDF (requires pip install cairosvg)
        try:
            from cairosvg import svg2pdf
            for svg_name, pdf_name in [("dfd.svg", "dfd.pdf"), ("logic.svg", "logic.pdf"), ("erd.svg", "erd.pdf")]:
                svg_path = os.path.join(output_dir, svg_name)
                if os.path.exists(svg_path):
                    pdf_path = os.path.join(output_dir, pdf_name)
                    svg2pdf(url=svg_path, write_to=pdf_path)
                    st.success(f"📄 PDF saved: {pdf_path}")
        except ImportError:
            st.warning("For PDF export, install cairosvg: pip install cairosvg")
        except Exception as e:
            st.error(f"PDF conversion failed: {e}")
        
        st.info(f"💾 Saved .mmd codes + SVGs/PDFs to {output_dir}/")
        
        output = f"""
### Extracted JSON:
```{extracted}```

### DFD:
```{dfd}```

### Logic Diagram:
```{logic}```

### DB ERD:
```{db}```
        """
        status.text("✅ Pipeline complete!")

    progress_bar.progress(1.0)
    st.balloons()  # Party confetti! 🎉

    # Display Results
    st.markdown("## 📊 Generated Diagrams")
    st.markdown(output)

    # Download Buttons (For Codes/Images)
    st.markdown("---")
    st.subheader("💾 Downloads")
    col_d1, col_d2, col_d3 = st.columns(3)
    
    with col_d1:
        st.download_button(
            "📥 Download DFD Code",
            dfd if 'dfd' in locals() else "No DFD generated",
            file_name="dfd.mmd",
            mime="text/plain"
        )
        # NEW: PDF Download
        if os.path.exists("output/dfd.pdf"):
            with open("output/dfd.pdf", "rb") as f:
                st.download_button(
                    "📥 Download DFD PDF",
                    f.read(),
                    file_name="dfd.pdf",
                    mime="application/pdf"
                )
    
    with col_d2:
        st.download_button(
            "📥 Download Logic Code",
            logic if 'logic' in locals() else "No Logic generated",
            file_name="logic.mmd",
            mime="text/plain"
        )
        # NEW: PDF Download
        if os.path.exists("output/logic.pdf"):
            with open("output/logic.pdf", "rb") as f:
                st.download_button(
                    "📥 Download Logic PDF",
                    f.read(),
                    file_name="logic.pdf",
                    mime="application/pdf"
                )
    
    with col_d3:
        st.download_button(
            "📥 Download ERD Code",
            db if 'db' in locals() else "No ERD generated",
            file_name="erd.mmd",
            mime="text/plain"
        )
        # NEW: PDF Download
        if os.path.exists("output/erd.pdf"):
            with open("output/erd.pdf", "rb") as f:
                st.download_button(
                    "📥 Download ERD PDF",
                    f.read(),
                    file_name="erd.pdf",
                    mime="application/pdf"
                )
    
    # Image Previews (If Rendered)
    if ag.MERMAID_AVAILABLE:
        st.markdown("## 🖼️ Rendered Images")
        col_i1, col_i2, col_i3 = st.columns(3)
        
        with col_i1:
            if os.path.exists("output/dfd.svg"):
                st.image("output/dfd.svg", caption="DFD", use_column_width=True)
        
        with col_i2:
            if os.path.exists("output/logic.svg"):
                st.image("output/logic.svg", caption="Logic", use_column_width=True)
        
        with col_i3:
            if os.path.exists("output/erd.svg"):
                st.image("output/erd.svg", caption="ERD", use_column_width=True)

else:
    st.info("👆 Upload a file or paste text, then hit Generate!")

# Footer
st.markdown("---")
st.markdown("*Built with ❤️ using Streamlit + LangChain. Questions? Check console.*")